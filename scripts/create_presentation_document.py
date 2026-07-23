"""Create the Turkish project-report Word document used for presentations."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path("docs/CMS-RAG-Proje-Sunum-Dokumani.docx")
NAVY = "123047"
TEAL = "007C83"
LIGHT = "EAF2F5"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    props.append(element)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title(document, title, subtitle=None):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Title"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    if subtitle:
        paragraph = document.add_paragraph(subtitle)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].font.color.rgb = RGBColor.from_string(TEAL)


def add_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else TEAL)
    return heading


def bullet(document, text, level=0):
    paragraph = document.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.add_run(text)


def numbered(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text)


def table(document, headers, rows, widths=None):
    result = document.add_table(rows=1, cols=len(headers))
    result.style = "Light Shading Accent 1"
    for idx, header in enumerate(headers):
        shade(result.rows[0].cells[idx], NAVY)
        set_cell_text(result.rows[0].cells[idx], header, bold=True, color="FFFFFF")
    for row in rows:
        cells = result.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
            if len(result.rows) % 2 == 0:
                shade(cells[idx], LIGHT)
    if widths:
        for row in result.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    document.add_paragraph()
    return result


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CMS-RAG Asistanı | Proje Sunum Dokümanı | Temmuz 2026")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string("5C6B73")


def build():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    add_footer(section)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(30)

    document.add_paragraph()
    document.add_paragraph()
    add_title(document, "CMS-RAG Asistanı", "Savaş Yönetim Sistemi Bilgi Erişim ve Karar Destek Asistanı")
    document.add_paragraph()
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run("Proje Sunumu İçin Teknik ve Yönetimsel Dokümantasyon\n").bold = True
    intro.add_run("Sürüm 1.0 | Temmuz 2026")
    document.add_paragraph()
    table(document, ["Proje", "Kapsam", "Çalışma Modu"], [["CMS-RAG Asistanı", "ADVENT CMS ve kamuya açık CMS referansları", "Yerel / on-premise"]])
    document.add_page_break()

    add_heading(document, "Yönetici Özeti")
    document.add_paragraph(
        "CMS-RAG Asistanı; Savaş Yönetim Sistemi (Combat Management System – CMS) "
        "dokümanlarını doğal dille, kaynak göstererek sorgulanabilir hâle getiren yerel "
        "bir yapay zekâ asistanıdır. Çözüm, büyük dil modelinin genel bilgisini doğrudan "
        "kullanmak yerine önce onaylı dokümanlardan kanıt toplar, ardından yanıt üretir."
    )
    bullet(document, "Resmî HAVELSAN/ADVENT içeriği, açık-kamu referanslarından fiziksel olarak ayrı indekslenir.")
    bullet(document, "Hibrit arama, anahtar kelime araması ile anlamsal aramayı birleştirir.")
    bullet(document, "Cross-encoder reranking, en alakalı parçaları LLM bağlamına taşır.")
    bullet(document, "Yanıt ile birlikte belge, sayfa ve yeniden sıralama skoru sunulur.")
    bullet(document, "Ollama ve yerel modeller ile hassas verilerin haricî LLM servislerine gönderilmesi engellenir.")

    add_heading(document, "1. Problem Tanımı ve Hedef", 1)
    document.add_paragraph(
        "CMS ekosisteminde ürün broşürleri, teknik kılavuzlar, standartlar ve operasyonel "
        "referanslar çoğunlukla farklı formatlarda ve farklı kaynaklarda bulunur. Bu durum "
        "bilgiye erişim süresini uzatır; ayrıca güvenilmeyen yanıt üretimi, savunma alanında "
        "kabul edilemez bir risk oluşturur."
    )
    table(document, ["Mevcut Zorluk", "Çözüm Yaklaşımı", "Beklenen Kazanım"], [
        ["Dağınık dokümanlar", "Tekil bilgi erişim katmanı", "Hızlı erişim"],
        ["Yanlış / kaynaksız yanıt riski", "RAG ve zorunlu kaynak görünürlüğü", "Denetlenebilirlik"],
        ["Resmî ve genel bilginin karışması", "Ayrı koleksiyonlar ve kapsam seçimi", "Güven sınırı"],
        ["Terim ve ifade çeşitliliği", "Dense + BM25 hibrit arama", "Daha yüksek erişim kalitesi"],
    ])

    add_heading(document, "2. Çözümün Değer Önerisi", 1)
    numbered(document, "Operatör, mühendis veya sunum hazırlayan kullanıcı; CMS konusunda sorusunu doğal dilde sorar.")
    numbered(document, "Sistem sorgu kapsamına göre resmî veya kamuya açık bilgi koleksiyonunu seçer.")
    numbered(document, "Hibrit erişim ve reranking sonrasında sadece ilgili kanıt parçaları modele iletilir.")
    numbered(document, "Yerel LLM teknik, kısa ve kaynak görünürlüğü olan yanıtı üretir.")

    add_heading(document, "3. Kurumsal Mimari", 1)
    document.add_paragraph("Sistem iki güven seviyeli bilgi koleksiyonu ve tek sorgu orkestratörü etrafında tasarlanmıştır.")
    table(document, ["Katman", "Bileşen", "Sorumluluk"], [
        ["Kullanıcı", "Streamlit arayüzü", "Sorgu kapsamı, PDF yükleme, kaynak görüntüleme"],
        ["Orkestrasyon", "CMSKnowledgeBase", "Koleksiyon seçimi, birleştirme, tekilleştirme"],
        ["Erişim", "FAISS + BM25", "Anlamsal ve anahtar kelime tabanlı aday bulma"],
        ["Kalite", "BAAI bge-reranker", "Adayların sorguya göre yeniden sıralanması"],
        ["Üretim", "Ollama / qwen2.5:3b", "Yerel yanıt üretimi"],
        ["Veri", "PDF, Markdown, kaynak kataloğu", "Provenance ve koleksiyon ayrımı"],
    ])
    document.add_paragraph("Sorgu akışı", style="Caption")
    flow = document.add_table(rows=1, cols=7)
    labels = ["Soru", "Kapsam", "Hybrid\nSearch", "Birleştir", "Rerank", "Yerel LLM", "Yanıt +\nKaynak"]
    for idx, label in enumerate(labels):
        shade(flow.rows[0].cells[idx], TEAL if idx in {2, 4} else NAVY)
        set_cell_text(flow.rows[0].cells[idx], label, bold=True, color="FFFFFF")
        flow.rows[0].cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    add_heading(document, "4. RAG Boru Hattı", 1)
    table(document, ["Aşama", "Uygulama", "Neden Önemli?"], [
        ["Toplama", "Onaylı web kataloğu + yüklenen PDF", "Kontrolsüz crawl yerine izlenebilir kaynak"],
        ["Yükleme", "PDF ve Markdown okuyucu", "Sayfa ve kaynak metaverisinin korunması"],
        ["Parçalama", "900 karakter, 150 karakter bindirme", "Bağlamı koruyan küçük erişim birimleri"],
        ["Embedding", "BAAI/bge-small-en-v1.5", "Anlamsal benzerlik"],
        ["Hibrit erişim", "FAISS dense + BM25", "Terim eşleşmesi ve kavramsal erişim"],
        ["Reranking", "BAAI/bge-reranker-base", "En iyi bağlamı seçme"],
        ["Yanıt", "qwen2.5:3b / Ollama", "Yerel, kontrollü üretim"],
    ])

    add_heading(document, "5. Veri Kaynakları ve Yönetişim", 1)
    document.add_paragraph("Kaynakların niteliği yanıtın güvenilirliğini belirler. Bu nedenle kaynak kataloğu, veri mimarisinin bir parçasıdır.")
    table(document, ["Koleksiyon", "İçerik", "Kural"], [
        ["havelsan", "Resmî ADVENT broşürü ve yetkili kurumsal dokümanlar", "Resmî ürün bilgisi için birincil kaynak"],
        ["open_source", "NATO interoperability ve NISP kamu referansları", "Genel kavram/standart bağlamı"],
        ["uploaded", "Kullanıcının yüklediği yerel PDF'ler", "Kurum onayı ve veri sınıflandırması gerekir"],
    ])
    bullet(document, "Her parça: belge adı, kaynak yolu, koleksiyon, otorite düzeyi ve sayfa numarası taşır.")
    bullet(document, "Kısıtlı dağıtımlı STANAG/MIL-STD belgeleri, lisans ve yetki doğrulanmadan sisteme alınmaz.")
    bullet(document, "Web senkronizasyonu yalnızca data/source_catalog.json içinde açıkça onaylanmış HTTPS adreslerine yapılır.")

    add_heading(document, "6. Güvenlik, Gizlilik ve Denetlenebilirlik", 1)
    table(document, ["Risk", "Mevcut Kontrol", "Üretim Ortamı İçin İlave Önlem"], [
        ["Veri sızıntısı", "Yerel Ollama ve local_files_only model yükleme", "Ağ segmentasyonu, şifreli disk"],
        ["Kaynak karışması", "Ayrı FAISS koleksiyonları", "Rol tabanlı erişim ve veri sınıflandırması"],
        ["Kaynak güvenilirliği", "Katalog ve metaveri", "Belge onay iş akışı, periyodik gözden geçirme"],
        ["İndeks bütünlüğü", "Yerel kontrol edilen indeks dizini", "İmzalı artefact, yedekleme, erişim günlüğü"],
        ["Halüsinasyon", "Kanıt odaklı prompt ve kaynak paneli", "Yanıt kalite eşiği, insan onayı"],
    ])

    add_heading(document, "7. Uygulama Demonstrasyonu", 1)
    document.add_paragraph("Sunum esnasında aşağıdaki demo hikâyesi önerilir:")
    numbered(document, "Uygulamayı açın ve 'Yalnızca HAVELSAN resmî kaynakları' kapsamını seçin.")
    numbered(document, "Sorguyu sorun: 'ADVENT hangi taktik veri linklerini destekler?'")
    numbered(document, "Yanıtta Link 11/16/22, SIMPLE, JREAP ve VMF bilgisini; kaynak panelinde ise broşür adı ve sayfa bilgisini gösterin.")
    numbered(document, "Kapsamı 'Birleşik' yapıp NATO birlikte çalışabilirlik kavramına dair bir takip sorusu sorun.")
    numbered(document, "Resmî ürün bilgisi ile genel standart bağlamının ayrı kaynaklardan geldiğini vurgulayın.")

    add_heading(document, "8. Mevcut Doğrulama Sonuçları", 1)
    table(document, ["Kontrol", "Sonuç", "Açıklama"], [
        ["Birim testi", "Başarılı", "Metin yükleyicisinin koleksiyon ve otorite metaverisi"],
        ["İndeks oluşturma", "Başarılı", "2 koleksiyon, toplam 68 parça"],
        ["Uçtan uca soru", "Başarılı", "HAVELSAN kapsamındaki ADVENT taktik veri linki sorgusu"],
        ["Arayüz açılışı", "Başarılı", "Streamlit yerel sunucusu doğrulandı"],
    ])
    document.add_paragraph("Not: Reranker skoru bir güven yüzdesi değildir; yalnızca adayların göreli alaka sırasını ifade eder.")

    add_heading(document, "9. Sınırlılıklar ve Gelişim Yol Haritası", 1)
    table(document, ["Dönem", "Önerilen İş", "Başarı Kriteri"], [
        ["Kısa vade", "Yanıt-kaynak bağını otomatik test eden değerlendirme seti", "Citation precision / recall"],
        ["Kısa vade", "Belge yaşam döngüsü ve artımlı indeksleme", "Yalnızca değişen belgenin işlenmesi"],
        ["Orta vade", "Yetkilendirme, audit log ve kullanıcı rolleri", "Kurumsal güvenlik gereksinimleri"],
        ["Orta vade", "Türkçe-İngilizce domain embedding karşılaştırması", "Ölçülmüş retrieval iyileşmesi"],
        ["Uzun vade", "Geri bildirim döngüsü ve insan denetimli kalite kontrol", "Kabul edilen yanıt kalitesi"],
    ])

    add_heading(document, "10. PowerPoint Sunum Önerisi", 1)
    table(document, ["Slayt", "Başlık", "Ana Mesaj"], [
        ["1", "CMS-RAG Asistanı", "Güvenilir bilgi erişimi için yerel yapay zekâ"],
        ["2", "Problem", "Dağınık dokümanlar ve kaynaksız yanıt riski"],
        ["3", "Çözüm", "Kanıt topla, yeniden sırala, yerelde yanıt üret"],
        ["4", "Mimari", "İki koleksiyon + hibrit arama + reranking"],
        ["5", "Güvenlik", "On-premise, kaynak ayrımı, yönetişim"],
        ["6", "Canlı Demo", "ADVENT veri linkleri sorusu ve kaynak paneli"],
        ["7", "Sonuç", "Denetlenebilir, genişletilebilir ve kurumsal temelli"],
    ])

    add_heading(document, "11. Kaynakça", 1)
    references = [
        "HAVELSAN, “ADVENT Network Supported Data Integrated Combat Management System”, resmî ürün sayfası: https://www.havelsan.com/en/solutions/advent-combat-management-system (erişim: Temmuz 2026).",
        "HAVELSAN, “ADVENT CMS Brochure”, yerel RAG koleksiyonu: data/raw/havelsan/havelsan-advent-brochure.pdf.",
        "NATO, “Interoperability: connecting forces”, https://www.nato.int/en/what-we-do/deterrence-and-defence/interoperability-connecting-forces (erişim: Temmuz 2026).",
        "NATO Interoperability Standards and Profiles (NISP), https://live.nisp.nw3.dk/ (erişim: Temmuz 2026).",
    ]
    for reference in references:
        bullet(document, reference)

    document.add_paragraph()
    closing = document.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run("Özet mesaj: Doğru bilgiye, doğru kaynakla ve doğru güven sınırı içinde erişim.")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
